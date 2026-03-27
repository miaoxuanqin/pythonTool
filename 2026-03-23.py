import os
import re
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def insert_paragraph_after(paragraph, text=None):
    """在指定段落后插入新段落"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    if text:
        # 处理txt中的段落逻辑：txt里的双换行视为新段落，单换行视为行内换行
        # 这里采用最严谨的方式：txt里的每一行都创建一个完整的Word段落
        return new_para  # 仅返回对象，由调用处处理内容
    return new_para


def copy_paragraph_format(src_para, tgt_para):
    """
    深度克隆段落格式：包括对齐方式、缩进、行间距、段前后距离等
    """
    tgt_para.alignment = src_para.alignment
    target_fmt = tgt_para.paragraph_format
    source_fmt = src_para.paragraph_format

    target_fmt.first_line_indent = source_fmt.first_line_indent
    target_fmt.line_spacing = source_fmt.line_spacing
    target_fmt.line_spacing_rule = source_fmt.line_spacing_rule
    target_fmt.space_before = source_fmt.space_before
    target_fmt.space_after = source_fmt.space_after
    target_fmt.left_indent = source_fmt.left_indent
    target_fmt.right_indent = source_fmt.right_indent
    target_fmt.keep_together = source_fmt.keep_together
    target_fmt.keep_with_next = source_fmt.keep_with_next


def set_run_format(run, sample_run):
    """克隆字符级格式：字体、字号、颜色等"""
    run.font.size = sample_run.font.size
    run.font.name = sample_run.font.name
    # 复制中文字体设置
    rPr_src = sample_run._element.get_or_add_rPr()
    rPr_tgt = run._element.get_or_add_rPr()

    # 复制字体映射（宋体等）
    src_fonts = rPr_src.find(qn('w:rFonts'))
    if src_fonts is not None:
        new_fonts = copy.deepcopy(src_fonts)
        rPr_tgt.append(new_fonts)


def fill_word_content(txt_path, word_path, output_path):
    # 1. 解析 TXT
    content_map = {}
    title_re = re.compile(r'^#+\s+([\d\.]+)\s+(.*)')
    current_title = None
    current_text_blocks = []  # 存储段落列表

    with open(txt_path, 'r', encoding='utf-8') as f:
        for line in f:
            match = title_re.match(line.strip())
            if match:
                if current_title:
                    content_map[current_title] = current_text_blocks
                current_title = match.group(2).strip()
                current_text_blocks = []
            else:
                if current_title and line.strip():
                    current_text_blocks.append(line.strip())
        if current_title:
            content_map[current_title] = current_text_blocks

    # 2. 打开 Word 并寻找“样板段落”
    doc = Document(word_path)

    # 寻找一个看起来像正文的段落作为样板（非标题、非空、长度适中）
    sample_para = None
    for p in doc.paragraphs:
        if 10 < len(p.text) < 100 and p.style.name == 'Normal':
            sample_para = p
            break
    if not sample_para:  # 兜底方案：取第一个普通段落
        sample_para = doc.paragraphs[10]

        # 3. 开始填充
    matched_count = 0
    # 注意：由于在遍历时修改文档会导致索引混乱，我们先记录位置
    targets = []
    for i, para in enumerate(doc.paragraphs):
        clean_para_text = para.text.replace(' ', '').replace('\t', '')
        for title_name, text_list in content_map.items():
            if title_name.replace(' ', '') in clean_para_text:
                targets.append((para, text_list))
                break

    for anchor_para, text_list in reversed(targets):
        # 倒序插入，保证位置不偏移
        # 对于txt里的每一行，都创建一个严格格式化的段落
        last_p = anchor_para
        for text in reversed(text_list):
            new_p = insert_paragraph_after(last_p)
            new_p.text = text

            # 复制样板格式
            copy_paragraph_format(sample_para, new_p)

            # 复制字体格式
            if sample_para.runs:
                for run in new_p.runs:
                    set_run_format(run, sample_para.runs[0])

            matched_count += 1

    doc.save(output_path)
    print(f"处理完成！通过采样‘样板段落’，严格同步了 {matched_count} 个段落的格式。")


# 路径
txt_file = r'D:\DESKTOP\txt.txt'
word_file = r'D:\DESKTOP\分工.docx'
output_file = r'D:\DESKTOP\分工2_严格格式填充.docx'

if __name__ == "__main__":
    fill_word_content(txt_file, word_file, output_file)