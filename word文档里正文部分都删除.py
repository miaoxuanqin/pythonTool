from docx import Document


def clear_only_body_text(file_path, save_path):
    try:
        doc = Document(file_path)

        # 遍历所有段落
        # 注意：由于删除列表元素会导致索引混乱，我们采用倒序删除或直接操作底层元素
        for paragraph in list(doc.paragraphs):
            # 获取该段落的样式名称
            style_name = paragraph.style.name

            # 如果样式是 'Normal' (普通正文) 或不包含 'Heading' (标题)
            # 你可以根据需要调整这个判断条件
            if "Heading" not in style_name:
                p = paragraph._element
                p.getparent().remove(p)
                paragraph._p = paragraph._element = None

        # 如果需要保留表格，就注释掉下面这两行；如果表格也要删，就保留
        # for table in doc.tables:
        #     t = table._element
        #     t.getparent().remove(t)

        doc.save(save_path)
        print(f"处理完成！保留了标题，正文已清空。")

    except Exception as e:
        print(f"运行出错: {e}")


input_path = r"D:\DESKTOP\模板 目录 全部.docx"
output_path = r"D:\DESKTOP\模板 目录 全部_仅留标题.docx"

if __name__ == "__main__":
    clear_only_body_text(input_path, output_path)