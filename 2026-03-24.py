import os
import re
from docx import Document
from docx.shared import Pt, Chars
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_paragraph_after(paragraph, text=None):
    """在指定段落后插入新段落并返回"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    if text:
        # 处理换行符，将 txt 中的 \n 转换为 Word 的换行
        lines = text.split('\n')
        for i, line in enumerate(lines):
            run = new_para.add_run(line.strip())
            if i < len(lines) - 1:
                run.add_break()
    return new_para

def set_para_format(paragraph):
    """强制设置严格的正文格式：宋体、小四、首行缩进2字符、单倍行距"""
    # 1. 段落整行设置
    paragraph.paragraph_format.first_line_indent = Chars(2) # 严格首行缩进2字符
    paragraph.paragraph_format.line_spacing = 1.0           # 单倍行距
    paragraph.paragraph_format.space_before = Pt(0)        # 段前0磅
    paragraph.paragraph_format.space_after = Pt(0)         # 段后0磅
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT          # 左对齐

    # 2. 字体详细设置
    for run in paragraph.runs:
        run.font.size = Pt(12)                             # 小四字号
        run.font.name = 'Times New Roman'                  # 西文字体
        # 强制设置中文字体为宋体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def fill_word_content(txt_path, word_path, output_path):
    # 解析 txt
    content_map = {}
    title_re = re.compile(r'^#+\s+([\d\.]+)\s+(.*)')
    current_title_name = None
    current_text = []

    with open(txt_path, 'r', encoding='utf-8') as f:
        for line in f:
            clean_line = line.strip()
            match = title_re.match(clean_line)
            if match:
                if current_title_name:
                    content_map[current_title_name] = "\n".join(current_text).strip()
                current_title_name = match.group(2).strip()
                current_text = []
            else:
                if current_title_name and clean_line:
                    current_text.append(clean_line)
        if current_title_name:
            content_map[current_title_name] = "\n".join(current_text).strip()

    # 打开 Word 并匹配
    doc = Document(word_path)
    matched_count = 0

    for para in doc.paragraphs:
        word_text = para.text.replace('\t', '').replace(' ', '').strip()
        if not word_text: continue

        for title_name, body_text in content_map.items():
            clean_title_name = title_name.replace(' ', '')
            if clean_title_name in word_text:
                # 插入并美化格式
                new_para = insert_paragraph_after(para, body_text)
                set_para_format(new_para)
                print(f"已严格格式化填充: {title_name}")
                matched_count += 1
                break

    doc.save(output_path)
    print(f"\n处理完成，共填充 {matched_count} 处。请查看：{output_path}")

# 路径设置
txt_file = r'D:\DESKTOP\txt.txt'
word_file = r'D:\DESKTOP\分工2.docx'
output_file = r'D:\DESKTOP\分工2_填充完成.docx'

if __name__ == "__main__":
    fill_word_content(txt_file, word_file, output_file)