import pandas as pd
from docx import Document
import os

def process_excel_to_word(excel_path, word_path, output_path):
    # 1. 读取 Excel 数据
    # 使用 openpyxl 引擎以支持 xlsx 格式
    df = pd.read_excel(excel_path, sheet_name='Sheet1')

    # 2. 处理合并单元格逻辑：向下填充 (Forward Fill)
    # 因为合并单元格在读取时，只有首行有值，其余行是 NaN
    df['系统'] = df['系统'].ffill()
    df['功能模块'] = df['功能模块'].ffill()
    # 功能点通常每行都有，但为了保险也处理一下
    df['功能点'] = df['功能点'].ffill()

    # 3. 加载或创建 Word 文档
    if os.path.exists(word_path):
        doc = Document(word_path)
    else:
        doc = Document()
        print(f"提示：未找到原始Word文件，将创建新文档。")

    # 用于记录上一次写入的内容，避免重复写入高层级标题
    last_system = None
    last_module = None

    # 4. 遍历数据并写入标题
    for index, row in df.iterrows():
        current_system = str(row['系统']).strip()
        current_module = str(row['功能模块']).strip()
        current_function = str(row['功能点']).strip()

        # 写入 6 级标题 (系统) - 只有当系统名称变化时才写入
        if current_system != last_system:
            doc.add_heading(current_system, level=6)
            last_system = current_system
            last_module = None # 系统变了，模块记录重置

        # 写入 7 级标题 (功能模块) - 只有当模块名称变化时才写入
        if current_module != last_module:
            doc.add_heading(current_module, level=7)
            last_module = current_module

        # 写入 8 级标题 (功能点) - 每一行都要写入
        doc.add_heading(current_function, level=8)

    # 5. 保存结果
    doc.save(output_path)
    print(f"处理完成！文件已保存至: {output_path}")

# --- 配置路径 ---
# 注意：Windows 路径建议使用双反斜杠 \\ 或在字符串前加 r
excel_file = r"D:\DESKTOP\功能清单分工.xlsx"
word_file = r"D:\DESKTOP\word.docx"
output_file = r"D:\DESKTOP\word_updated.docx"

if __name__ == "__main__":
    try:
        process_excel_to_word(excel_file, word_file, output_file)
    except Exception as e:
        print(f"运行出错: {e}")