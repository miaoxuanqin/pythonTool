import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def get_word_numbering(paragraph):
    """
    尝试获取 Word 段落的自动编号字符串。
    注意：此函数依赖于 Word 文档中编号列表的 XML 结构。
    """
    # 这是一个简化处理，因为直接从 XML 提取完整编号字符串非常复杂
    # 如果 Word 是标准的自动编号，Python 读取 text 只有文字
    # 如果您的 Word 编号是手动输入的，text 会包含编号
    return paragraph.text.strip()


def solve_task():
    txt_path = r'D:\DESKTOP\txt.txt'
    docx_path = r'D:\DESKTOP\目录 模板 苗 内容空白.docx'
    output_path = r'D:\DESKTOP\更新后的文档_双重锁定版.docx'

    if not os.path.exists(txt_path) or not os.path.exists(docx_path):
        print("错误：路径不存在。")
        return

    # 1. 解析 TXT：锁定 (编号, 文字) 作为唯一 Key
    # 数据结构: { ("2.4.1.3.1.", "标题文字"): ["正文行1", "正文行2"] }
    txt_data = {}
    current_key = None
    current_content = []

    try:
        with open(txt_path, 'r', encoding='utf-8-sig') as f:
            lines = f.readlines()
    except:
        with open(txt_path, 'r', encoding='gbk') as f:
            lines = f.readlines()

    for line in lines:
        line_strip = line.strip()
        # 匹配 # 2.4.1.3.1.标题文字 或 # 2.4.1.3.1. 标题文字
        match = re.match(r'^#+\s+([\d\.]+)\s*(.*)', line_strip)
        if match:
            if current_key:
                txt_data[current_key] = current_content

            number_part = match.group(1).strip()
            text_part = match.group(2).strip()
            current_key = (number_part, text_part)  # 双重锁定 Key
            current_content = []
        else:
            if current_key and line_strip:
                current_content.append(line_strip)
    if current_key:
        txt_data[current_key] = current_content

    # 2. 操作 Word
    doc = Document(docx_path)

    # 遍历 Word 段落进行匹配
    for para in doc.paragraphs:
        word_full_text = para.text.strip()

        # 关键逻辑：
        # 如果您的 Word 编号是“自动编号”，para.text 通常只含文字。
        # 如果您的 Word 编号是“手动输入”，para.text 会包含编号。

        target_found = False
        target_content = []

        for (txt_num, txt_title), content in txt_data.items():
            # 匹配策略：
            # 情况A：Word text 包含了编号和文字 (例如 "2.4.1.3.1.国企央企...")
            # 情况B：Word 自动编号不可见，但文字匹配，且我们根据上下文逻辑校验
            if txt_num in word_full_text and txt_title in word_full_text:
                target_found = True
                target_content = content
                break

        if target_found:
            print(f"双重锁定成功: {txt_num} {txt_title}")

            # 倒序插入以保持 TXT 原有行序
            for text_line in reversed(target_content):
                new_p_element = doc.add_paragraph()._element
                para._element.addnext(new_p_element)
                p_obj = Paragraph(new_p_element, doc)

                # 格式设置：首行缩进 2 字符 (小四12pt -> 24pt)
                p_obj.paragraph_format.first_line_indent = Pt(24)

                run = p_obj.add_run(text_line)
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)

    # 3. 保存
    doc.save(output_path)
    print(f"\n任务完成！文件保存在: {output_path}")


if __name__ == "__main__":
    solve_task()