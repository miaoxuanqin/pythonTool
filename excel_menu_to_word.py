import os
from openpyxl import load_workbook
from docx import Document


def build_merged_cell_map(ws):
    """预处理合并单元格映射，解决3600行数据读取性能问题"""
    merged_map = {}
    if not ws.merged_cells:
        return merged_map
    for merged_range in ws.merged_cells.ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        top_left_value = ws.cell(row=min_row, column=min_col).value
        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                merged_map[(r, c)] = top_left_value
    return merged_map


def process_excel_to_word(excel_path, word_path, output_path):
    print("正在加载 Excel...")
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    merged_map = build_merged_cell_map(ws)

    print("正在提取 Excel 数据树...")
    data_tree = []
    # 映射关系：B(2):5级, C(3):6级, D(4):7级, E(5):8级, F(6):9级, G(7):10级
    for r in range(2, ws.max_row + 1):
        row_content = []
        for c in range(2, 8):
            val = merged_map.get((r, c), ws.cell(row=r, column=c).value)
            row_content.append(str(val).strip() if val is not None else "")
        if any(row_content):
            data_tree.append(row_content)

    print("正在解析 Word 模板并捕获样板样式...")
    doc = Document(word_path)

    style_objects = {}  # 存储级别对应的样式对象
    placeholders_to_remove = []  # 存储需要删除的样板行
    anchor_para = None

    # 遍历文档寻找锚点和样板样式
    for para in doc.paragraphs:
        text = para.text.strip()
        if "平台对接（苗）" == text:
            anchor_para = para
        elif "5级标题" == text:
            style_objects[5] = para.style
            placeholders_to_remove.append(para)
        elif "6级标题" == text:
            style_objects[6] = para.style
            placeholders_to_remove.append(para)
        elif "7级标题" == text:
            style_objects[7] = para.style
            placeholders_to_remove.append(para)
        elif "8级标题" == text:
            style_objects[8] = para.style
            placeholders_to_remove.append(para)
        elif "9级标题" == text:
            style_objects[9] = para.style
            placeholders_to_remove.append(para)

    if not anchor_para:
        print("错误：未在文档中找到锚点 '功能设计'，请检查Word内容")
        return

    print("正在按照捕获的样式插入 Excel 内容...")
    last_inserted = [None] * 6
    insert_cursor = anchor_para

    for row in data_tree:
        for i, title_text in enumerate(row):
            if not title_text or title_text.lower() == "none" or title_text == "":
                continue

            # 去重逻辑：父级相同且内容相同时跳过，避免合并单元格重复触发
            if title_text == last_inserted[i] and (i == 0 or row[i - 1] == last_inserted[i - 1]):
                continue

            current_level = i + 5
            new_para = doc.add_paragraph(title_text)

            # 应用捕获到的样式
            if current_level in style_objects:
                new_para.style = style_objects[current_level]
            else:
                # 针对10级标题或超出范围的处理：继承9级样式并手动加粗
                new_para.style = style_objects.get(9, doc.styles['Normal'])
                if current_level > 9 and new_para.runs:
                    new_para.runs[0].bold = True

                    # 将新段落插入到当前光标之后
            insert_cursor._element.addnext(new_para._element)
            insert_cursor = new_para

            # 更新状态
            last_inserted[i] = title_text
            for j in range(i + 1, 6):
                last_inserted[j] = None

    print("正在清理模板中的样板占位符...")
    for p in placeholders_to_remove:
        p._element.getparent().remove(p._element)

    # 保存文件
    try:
        doc.save(output_path)
        print(f"处理成功！输出文件已保存至：{output_path}")
    except Exception as e:
        print(f"保存失败，请检查文件是否被占用：{e}")


if __name__ == "__main__":
    # 请根据实际环境确认路径
    config = {
        "excel": r"D:\DESKTOP\海易办及平台对接清单.xlsx",
        "word": r"D:\DESKTOP\模板.docx",
        "output": r"D:\DESKTOP\模板_已更新.docx"
    }

    if os.path.exists(config["excel"]) and os.path.exists(config["word"]):
        process_excel_to_word(config["excel"], config["word"], config["output"])
    else:
        print("错误：未找到输入文件，请检查 D:\DESKTOP\ 路径。")