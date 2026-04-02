import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


docx_file = r"D:\DESKTOP\目录 模板 苗 内容空白.docx"
txt_file = r"D:\DESKTOP\txt.txt"
output_file = r"D:\DESKTOP\填充后文档_最终版.docx"


# TXT_PATH = r'D:\DESKTOP\txt.txt'
# WORD_PATH = r'D:\DESKTOP\目录 模板 苗 内容空白.docx'
# OUTPUT_PATH = r'D:\DESKTOP\填充后文档_最终版.docx'


# -----------------------------
# 标题key
# -----------------------------
def normalize_key(num, title):

    title = re.sub(r"\s+", "", title)

    return num + title


# -----------------------------
# 解析TXT
# -----------------------------
def parse_txt(path):

    with open(path, "r", encoding="utf-8") as f:
        text = f.read()

    pattern = r"(\d+(?:\.\d+)+)\s*([^\n]+)"

    matches = list(re.finditer(pattern, text))

    data = {}

    for i, m in enumerate(matches):

        num = m.group(1)
        title = m.group(2)

        key = normalize_key(num, title)

        start = m.end()

        if i + 1 < len(matches):
            end = matches[i + 1].start()
        else:
            end = len(text)

        content = text[start:end].strip()

        paragraphs = []

        for p in re.split(r"\n\s*\n", content):

            p = p.strip()

            if p:
                paragraphs.append(p)

        data[key] = paragraphs

    return data


# -----------------------------
# 插入段落
# -----------------------------
def insert_after(paragraph):

    new_p = OxmlElement("w:p")

    paragraph._p.addnext(new_p)

    new_para = paragraph._parent.add_paragraph()

    new_para._p = new_p

    return new_para


# -----------------------------
# 正文格式
# -----------------------------
def format_body(para, text):

    para.text = text

    run = para.runs[0]

    run.font.name = "宋体"

    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    run.font.size = Pt(12)

    pf = para.paragraph_format

    pf.first_line_indent = Inches(0.28)

    pf.line_spacing = 1.5


# -----------------------------
# 生成Word编号
# -----------------------------
def build_word_index(doc):

    counters = [0] * 10

    titles = []

    for para in doc.paragraphs:

        style = para.style.name

        if not style.startswith("Heading"):
            continue

        level = int(style.split()[-1])

        counters[level] += 1

        for i in range(level + 1, 10):
            counters[i] = 0

        nums = []

        for i in range(1, level + 1):
            if counters[i] > 0:
                nums.append(str(counters[i]))

        num = ".".join(nums) + "."

        titles.append((para, num, para.text.strip()))

    return titles


# -----------------------------
# 填充Word
# -----------------------------
def fill_docx(docx_path, txt_data):

    doc = Document(docx_path)

    titles = build_word_index(doc)

    for para, num, title in titles:

        key = normalize_key(num, title)

        if key not in txt_data:
            continue

        current = para

        for p in txt_data[key]:

            new_para = insert_after(current)

            format_body(new_para, p)

            current = new_para

    if os.path.exists(output_file):
        os.remove(output_file)

    doc.save(output_file)


# -----------------------------
# main
# -----------------------------
def main():

    txt_data = parse_txt(txt_file)

    fill_docx(docx_file, txt_data)

    print("生成完成:", output_file)


if __name__ == "__main__":
    main()