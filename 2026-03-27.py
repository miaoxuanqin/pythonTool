import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ================= 配置路径 =================
TXT_PATH = r'D:\DESKTOP\txt.txt'
WORD_PATH = r'D:\DESKTOP\目录 模板 苗 内容空白.docx'
OUTPUT_PATH = r'D:\DESKTOP\填充后文档_最终版.docx'


def clean_text(text):
    """清洗字符串：去除所有空格和开头的数字编号"""
    if not text: return ""
    text = "".join(text.split())
    # 去除开头的编号部分 (例如 2.1.1.3...)
    text = re.sub(r'^[\d\.]+', '', text)
    return text


def parse_txt(file_path):
    """
    改进后的解析逻辑：
    按行读取，遇到任何以 # 开头的行都视为标题。
    """
    data_map = {}
    current_title = None
    current_body = []

    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line: continue

            # 只要是以 # 开头的，不管有几个 #，都是标题
            if line.startswith('#'):
                # 保存上一个标题的内容
                if current_title and current_body:
                    data_map[current_title] = current_body

                # 提取新标题：去掉开头的 # 号和空格
                current_title = line.lstrip('#').strip()
                current_body = []
            else:
                # 不是以 # 开头的行，才是真正的正文
                if current_title is not None:
                    current_body.append(line)

        # 保存最后一个块
        if current_title and current_body:
            data_map[current_title] = current_body

    return data_map


def apply_strict_style(paragraph):
    """格式设置：宋体、小四、1.5倍行距、首行缩进"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(24)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # 必须先添加 Run 才能设置字体，或者对已有 Run 设置
    if not paragraph.runs:
        paragraph.add_run()

    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def process_word():
    print("正在解析 TXT 文件（已修正标题识别逻辑）...")
    txt_data = parse_txt(TXT_PATH)
    print(f"解析完成，共提取到 {len(txt_data)} 个标题及其对应正文。")

    if not os.path.exists(WORD_PATH):
        print("错误：未找到 Word 模板文件。")
        return

    doc = Document(WORD_PATH)
    tasks = []

    # 预处理 TXT 标题用于匹配
    clean_txt_titles = {clean_text(k): (k, v) for k, v in txt_data.items()}

    # 1. 扫描 Word 定位标题
    for i, para in enumerate(doc.paragraphs):
        word_text = para.text.strip()
        if not word_text: continue

        c_word = clean_text(word_text)

        for c_key, (original_key, body_list) in clean_txt_titles.items():
            # 使用全文匹配或包含匹配
            if c_key and c_key == c_word:  # 建议用 == 提高精确度，防止误触
                tasks.append((i, body_list, original_key))
                break

    # 2. 逆序插入，防止索引崩溃和死循环
    tasks.sort(key=lambda x: x[0], reverse=True)

    print(f"开始填充正文...")
    for index, body_list, title_name in tasks:
        # 在标题段落的下一个位置开始插入
        target_pos = index + 1

        # 逆序插入 body_list 以保持正文原有顺序
        for text_content in reversed(body_list):
            if target_pos < len(doc.paragraphs):
                new_p = doc.paragraphs[target_pos].insert_paragraph_before(text_content)
            else:
                new_p = doc.add_paragraph(text_content)
            apply_strict_style(new_p)

        print(f"已完成: {title_name}")

    # 3. 保存文件
    try:
        doc.save(OUTPUT_PATH)
        print(f"\n任务成功！文件保存至: {OUTPUT_PATH}")
    except PermissionError:
        print(f"\n错误：保存失败，请关闭已打开的 {os.path.basename(OUTPUT_PATH)}")


if __name__ == "__main__":
    process_word()