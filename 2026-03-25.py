import pandas as pd
from docx import Document
import os


def append_excel_titles_to_word(excel_path, word_path, output_path):
    # 1. 读取 Excel 数据
    # 指定引擎为 openpyxl 处理 xlsx 文件
    try:
        df = pd.read_excel(excel_path)
    except Exception as e:
        print(f"读取 Excel 出错: {e}")
        return

    # 2. 处理合并单元格逻辑：使用前向填充 (ffill) 还原层级关系
    # 仅选取前三列（对应 3、4、5 级标题）
    df_filled = df.iloc[:, 0:3].ffill()

    # 3. 加载 Word 文档
    if os.path.exists(word_path):
        doc = Document(word_path)
    else:
        print(f"提示：未找到模板 {word_path}，将创建新文档。")
        doc = Document()

    # 记录上一行处理的状态，用于去重
    last_h3, last_h4, last_h5 = None, None, None

    # 4. 遍历并写入
    for _, row in df_filled.iterrows():
        # 使用 iloc 避免 FutureWarning
        val3 = str(row.iloc[0]).strip()
        val4 = str(row.iloc[1]).strip()
        val5 = str(row.iloc[2]).strip()

        # 定义一个过滤函数，排除空值或 Excel 的列标题
        def is_valid(v):
            return v not in ['nan', 'None', '', '3级', '4级', '5级']

        # --- 处理 3 级标题 ---
        if is_valid(val3) and val3 != last_h3:
            doc.add_paragraph(val3, style='Heading 3')
            last_h3 = val3
            last_h4, last_h5 = None, None  # 父级更新，重置子级记录

        # --- 处理 4 级标题 ---
        if is_valid(val4) and val4 != last_h4:
            doc.add_paragraph(val4, style='Heading 4')
            last_h4 = val4
            last_h5 = None  # 父级更新，重置子级记录

        # --- 处理 5 级标题 ---
        if is_valid(val5) and val5 != last_h5:
            doc.add_paragraph(val5, style='Heading 5')
            last_h5 = val5

    # 5. 保存结果
    try:
        doc.save(output_path)
        print(f"【处理成功】")
        print(f"输入文件：{excel_path}")
        print(f"输出文件：{output_path}")
    except PermissionError:
        print(f"【错误】无法保存文件，请检查 {output_path} 是否已被 Word 打开。")


# --- 配置本地路径 ---
config = {
    "excel": r'D:\DESKTOP\海易办及平台对接清单.xlsx',
    "template": r'D:\DESKTOP\模板.docx',
    "output": r'D:\DESKTOP\海易办及平台对接清单_追加版.docx'
}

if __name__ == "__main__":
    append_excel_titles_to_word(config["excel"], config["template"], config["output"])