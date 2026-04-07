import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn


def process_txt_to_word_formatted(txt_path, docx_path, output_path):
    if not os.path.exists(txt_path):
        print(f"错误: 找不到文件 {txt_path}")
        return
    if not os.path.exists(docx_path):
        print(f"错误: 找不到模板文件 {docx_path}")
        return

    # 加载 Word 模板
    doc = Document(docx_path)

    # 读取 TXT 内容
    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 匹配标题行：以 # 开头，提取编号后的文字
        header_match = re.match(r'^#*\s*([\d\.]+)\s*(.*)', line)

        if header_match:
            dots_sequence = header_match.group(1)
            # 通过分割点号来计算数字的个数
            level = len(dots_sequence.split('.'))
            clean_title_text = header_match.group(2).strip()

            if not clean_title_text:
                clean_title_text = "未命名标题"

            # 插入标题（自动去掉编号）
            if 1 <= level <= 9:
                doc.add_heading(clean_title_text, level=level)
            else:
                p = doc.add_paragraph()
                run = p.add_run(clean_title_text)
                run.bold = True
        else:
            # 插入正文段落
            p = doc.add_paragraph()

            # --- 设置首行缩进 2 字符 ---
            # 在中文字体下，1个字符通常等于字号大小。小四是12pt，2字符即24pt
            p.paragraph_format.first_line_indent = Pt(24)

            # --- 设置正文字体与字号 ---
            run = p.add_run(line)
            run.font.size = Pt(12)  # 小四字号

            # 设置中文字体（可选，确保宋体等生效）
            run.font.name = 'SimSun'
            r = run._element.rPr.rFonts
            r.set(qn('w:eastAsia'), 'SimSun')

    # 保存结果
    doc.save(output_path)
    print(f"处理完成！已应用小四字号及首行缩进，文件保存在: {output_path}")


# --- 配置路径 ---
txt_file = r'D:\DESKTOP\txt.txt'
template_file = r'D:\DESKTOP\目录.docx'
output_file = r'D:\DESKTOP\生成追加.docx'

# 执行
process_txt_to_word_formatted(txt_file, template_file, output_file)