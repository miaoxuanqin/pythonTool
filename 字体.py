from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def unify_format_smart_indent(input_path, output_path):
    doc = Document(input_path)

    # 小四对应 12 磅
    target_size_pt = 12
    xml_size_val = str(target_size_pt * 2)
    # 定义 2 个字符的缩进量 (12pt * 2 = 24pt)
    indent_size = Pt(24)

    for paragraph in doc.paragraphs:
        # --- 1. 统一设置字号 ---
        for run in paragraph.runs:
            run.font.size = Pt(target_size_pt)
            rPr = run._element.get_or_add_rPr()

            # 西文/数字字号
            sz = rPr.get_or_add_sz()
            sz.set(qn('w:val'), xml_size_val)

            # 中文字号 (处理 szCs)
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = rPr.makeelement(qn('w:szCs'))
                rPr.append(szCs)
            szCs.set(qn('w:val'), xml_size_val)

        # --- 2. 智能设置首行缩进 ---
        # 只有当段落完全没有设置首行缩进（为 None）或缩进为 0 时，才进行设置
        # 如果用户已经手动缩进了（无论缩进多少），都不再处理
        current_indent = paragraph.paragraph_format.first_line_indent

        if current_indent is None or current_indent == 0:
            # 此外，通常我们会过滤掉“居中”对齐的段落（如标题），避免标题也跟着缩进
            if paragraph.alignment is None or paragraph.alignment == 0:  # 0 代表左对齐
                paragraph.paragraph_format.first_line_indent = indent_size

    # --- 3. 处理表格（通常表格内不缩进，仅统一字号） ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(target_size_pt)
                        rPr = run._element.get_or_add_rPr()
                        sz = rPr.get_or_add_sz()
                        sz.set(qn('w:val'), xml_size_val)
                        szCs = rPr.find(qn('w:szCs'))
                        if szCs is None:
                            szCs = rPr.makeelement(qn('w:szCs'))
                            rPr.append(szCs)
                        szCs.set(qn('w:val'), xml_size_val)

    doc.save(output_path)
    print(f"处理完成！已跳过原有缩进，保存至: {output_path}")


# 执行
unify_format_smart_indent('D:\\DESKTOP\\海口市智能网联汽车“车路云一体化”应用试点（一期）项目投标文件-技术方案模版.docx', 'D:\\DESKTOP\\智能缩进后的文档.docx')
